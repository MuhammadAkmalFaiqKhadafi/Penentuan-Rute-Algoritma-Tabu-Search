from flask import Flask, render_template, request, flash, redirect, send_file
from cmvrp_tabu_search import (
    run_cmvrp, load_input_data, load_vehicle_count,
    allocate_customers, tabu_search_vrp, fuel_price, fuel_consumption, 
    calculate_route_segments, generate_leaflet_html, laporan_rute)
from osrm_cache import osrm_cached_request, osrm_request_for_report
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from base64 import b64decode
from io import BytesIO
from PIL import Image
import json, os
import pandas as pd
import subprocess
import socket
import time

def clean_temp_files_on_startup():
    files_to_delete = [
        "data_route_results.json",
        "vrp_tabu_search_visualisasi.html",
        "laporan_cmvrp.csv",
        "laporan_segment_cmvrp.csv",
        "templates/laporan_rute_output.html",
        "templates/laporan_profit_output.html",
        "osrm_cache.json"  # kalau kamu juga ingin menghapus cache saat restart
    ]

    for file in files_to_delete:
        try:
            if os.path.exists(file):
                os.remove(file)
                print(f"[AUTO CLEAN] File dihapus saat Flask start: {file}")
        except Exception as e:
            print(f"[AUTO CLEAN ERROR] Tidak bisa hapus {file}: {e}")

# Jalankan pembersihan saat Flask start
clean_temp_files_on_startup()

def is_osrm_running(host='localhost', port=5000):
    """Cek apakah OSRM service sudah aktif di port tertentu"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        return s.connect_ex((host, port)) == 0

def start_osrm():
    """Jalankan OSRM server menggunakan Docker jika belum aktif"""
    if not is_osrm_running():
        print("OSRM belum aktif. Menjalankan container Docker OSRM...")
        subprocess.Popen([
            "docker", "run", "-t", "-i", "-p", "5000:5000",
            "-v", f"{os.getcwd()}:/data",
            "osrm/osrm-backend",
            "osrm-routed",
            "--algorithm", "mld", "/data/kalimantan-latest.osrm"
        ])
        # Tunggu beberapa detik agar OSRM siap menerima request
        time.sleep(5)
    else:
        print("OSRM sudah aktif.")

def load_route_results():
    with open("data_route_results.json", "r") as f:
        return json.load(f)

app = Flask(__name__)
app.secret_key = "rahasia_anda"

@app.route('/')
def index():
    hasil_tersedia = os.path.exists("data_route_results.json")

    if os.path.exists("laporan_cmvrp.csv"):
        df = pd.read_csv("laporan_cmvrp.csv")
        kendaraan_aktif = df.shape[0]
        kendaraan_total = load_vehicle_count()
        kendaraan_nonaktif = kendaraan_total - kendaraan_aktif

        total_customer_served = sum(df["Revenue"] > 0)
        total_jarak = df["Jarak Tempuh (km)"].sum()
        total_biaya = df["Biaya"].sum()
        total_profit = df["Profit"].sum()
    else:
        kendaraan_total = load_vehicle_count()
        kendaraan_aktif = 0
        kendaraan_nonaktif = kendaraan_total
        total_customer_served = 0
        total_jarak = 0
        total_biaya = 0
        total_profit = 0
        flash("Belum ada laporan CMVRP. Silakan jalankan proses terlebih dahulu.", "warning")

    return render_template("index.html",
                           kendaraan_aktif=kendaraan_aktif,
                           kendaraan_nonaktif=kendaraan_nonaktif,
                           total_customer_served=total_customer_served,
                           total_jarak=round(total_jarak, 2),
                           total_biaya=round(total_biaya, 2),
                           total_profit=round(total_profit, 2),
                           hasil_tersedia=hasil_tersedia)

@app.route('/set_kendaraan', methods=['POST'])
def set_kendaraan():
    vehicle_count = request.form.get('vehicle_count', type=int)
    if vehicle_count:
        config = {"vehicle_count": vehicle_count}
        with open("config.json", "w") as f:
            json.dump(config, f)
    return redirect('/')

@app.route('/submit', methods=['POST'])
def submit():
    name = request.form['name']
    tipe = request.form['type']
    lat = float(request.form['lat'])
    lon = float(request.form['lon'])
    demand = request.form.get('demand', type=int)
    fee = request.form.get('fee', type=int)
    supply = request.form.get('supply', type=int)

    data = {
        "name": name,
        "type": tipe,
        "lat": lat,
        "lon": lon,
        "demand": demand,
        "fee": fee
    }

    if tipe == "depot":
        data["supply"] = supply
    
    if not os.path.exists("data_lokasi.json"):
        with open("data_lokasi.json", "w") as f:
            json.dump([], f)

    with open("data_lokasi.json", "r+") as f:
        lokasi = json.load(f)
        lokasi.append(data)
        f.seek(0)
        json.dump(lokasi, f, indent=4)

    return redirect('/lokasi')

@app.route('/hapus/<int:index>', methods=['POST'])
def hapus(index):
    if os.path.exists("data_lokasi.json"):
        with open("data_lokasi.json", "r+") as f:
            data = json.load(f)
            if 0 <= index < len(data):
                data.pop(index)
                f.seek(0)
                f.truncate()
                json.dump(data, f, indent=4)
    return redirect('/lokasi')

@app.route('/lokasi')
def lokasi():
    if os.path.exists("data_lokasi.json"):
        with open("data_lokasi.json", "r") as f:
            data = json.load(f)
    else:
        data = []

    total_supply = sum(d.get("supply", 0) for d in data if d["type"] == "depot")
    total_demand = sum(d.get("demand", 0) for d in data if d["type"] == "customer")
    cukup = total_supply >= total_demand
    hasil_tersedia = os.path.exists("data_route_results.json")


    return render_template('lokasi.html', data=data,
                           total_supply=total_supply,
                           total_demand=total_demand,
                           cukup=cukup, hasil_tersedia=hasil_tersedia)

@app.route('/visualisasi')
def visualisasi():
    # Cek apakah hasil proses sudah ada
    if not os.path.exists("data_route_results.json"):
        flash("Belum ada hasil visualisasi. Silakan jalankan proses terlebih dahulu.", "warning")
        return redirect("/")
    
    if not os.path.exists("data_route_results.json"):
        return "Silakan klik tombol Proses dulu."

    route_data = load_route_results()
    _, customers, rest_areas, menginap_locs = load_input_data()

    # Ubah format route_data agar sesuai format parameter untuk generate_leaflet_html
    formatted_routes = []
    for route in route_data:
        formatted_routes.append((
            route["route_coords"],
            route["assigned_customers"],
            route.get("overnight", 0),
            route.get("total_distance", 0),
            route.get("total_revenue", 0),
            route.get("total_cost", 0),
            route.get("total_profit", 0),
            route.get("service_time", 0),
            route.get("rest_time", 0),
            route.get("nginap_time", 0),
            route.get("total_time", 0),
            route.get("travel_time", 0),
            route.get("work_time", 0)
        ))

    generate_leaflet_html(formatted_routes, rest_areas, menginap_locs)
    if not os.path.exists("vrp_tabu_search_visualisasi.html"):
        flash("File visualisasi tidak ditemukan. Silakan jalankan proses terlebih dahulu.", "warning")
        return redirect("/")
    return send_file("vrp_tabu_search_visualisasi.html")

@app.route('/export_segment')
def export_segment():
    if not os.path.exists("data_route_results.json"):
        return "Silakan klik tombol Proses dulu."

    with open("data_route_results.json", "r") as f:
        route_data = json.load(f)

    _, customers, _, _ = load_input_data()

    segment_rows = []

    for vehicle in route_data:
        i = vehicle["vehicle_id"]
        assigned_customers = vehicle["assigned_customers"]
        depot_coord = vehicle["depot"][1]

        # Ambil nama pelanggan dari assigned_customers
        route_customer_names = [name for _, name, _ in assigned_customers]

        segments_info = calculate_route_segments(route_customer_names, customers, depot_coord)
        for idx, seg in enumerate(segments_info["segments"], 1):
            segment_rows.append({
                "Kendaraan": f"Kendaraan {i}",
                "Segment": idx,
                "Titik Asal": seg["from"],
                "Titik Tujuan": seg["to"],
                "Jarak (km)": round(seg["distance_km"], 2),
                "Durasi (menit)": round(seg["duration_hr"] * 60, 2),
                "Biaya (Rp)": round(seg["cost"], 2)
            })

    df = pd.DataFrame(segment_rows)
    df.to_csv("laporan_segment_cmvrp.csv", index=False)
    return send_file("laporan_segment_cmvrp.csv", as_attachment=True)

@app.route('/export_word')
def export_word():
    # Render ulang HTML jika file tidak ada
    if not os.path.exists("templates/laporan_profit_output.html"):
        with app.test_request_context():
            tampilkan_laporan()
    if not os.path.exists("templates/laporan_rute_output.html"):
        with app.test_request_context():
            tampilkan_laporan_rute()

    html_path = "templates/laporan_profit_output.html"
    rute_path = "templates/laporan_rute_output.html"
    chart_path = "static/chart.png"

    doc = Document()
    doc.add_heading("Laporan Rute dan Biaya", 0)

    # ---------------- RINGKASAN & PROFIT ----------------
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")

    # Cari semua table-bordered dan ambil yang mengandung baris <td>
    table_total_candidates = soup.find_all("table", class_="table-bordered")
    table_total = None
    for t in table_total_candidates:
        if t.find("td"):  # pastikan tabel ini berisi data
            table_total = t
            break
    if table_total:
        doc.add_heading("Ringkasan Hasil", level=1)
        rows = table_total.find_all("tr")
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Kategori"
        hdr[1].text = "Nilai"
        for row in rows:
            cols = row.find_all("td")
            if len(cols) == 2:
                cells = table.add_row().cells
                cells[0].text = cols[0].get_text(strip=True)
                cells[1].text = cols[1].get_text(strip=True)

    doc.add_paragraph()

    table_detail = soup.find("table", class_="table-striped")
    if table_detail:
        doc.add_heading("Rincian hasil tiap Kendaraan", level=1)
        headers = [th.get_text(strip=True) for th in table_detail.find_all("th")]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for row in table_detail.find_all("tr")[1:]:
            cols = row.find_all("td")
            if len(cols) == len(headers):
                cells = table.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = c.get_text(strip=True)

    # ---------------- DETAIL RUTE ----------------
    with open(rute_path, "r", encoding="utf-8") as f:
        soup_rute = BeautifulSoup(f, "html.parser")

    doc.add_heading("Optimasi Rute Setiap Kendaraan", level=1)

    cards = soup_rute.find_all("div", class_="card")
    for card in cards:
        header = card.find("div", class_="card-header")
        if header:
            doc.add_heading(header.text.strip(), level=2)

        table_summary = card.find("table", class_="table-bordered")
        if table_summary:
            headers = [th.get_text(strip=True) for th in table_summary.find_all("th")]
            values = [td.get_text(strip=True) for td in table_summary.find_all("td")]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            row_cells = table.add_row().cells
            for i, v in enumerate(values):
                row_cells[i].text = v

        table_detail = card.find("table", class_="table-striped")
        if table_detail:
            headers = [th.get_text(strip=True) for th in table_detail.find_all("th")]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
            for tr in table_detail.find_all("tr")[1:]:
                tds = tr.find_all("td")
                if len(tds) == len(headers):
                    row_cells = table.add_row().cells
                    for i, td in enumerate(tds):
                        row_cells[i].text = td.get_text(strip=True)

        doc.add_paragraph()

    # ---------------- GRAFIK CHART ----------------
    if os.path.exists(chart_path):
        doc.add_heading("Grafik", level=1)
        doc.add_picture(chart_path, width=Inches(5.5))
        print("[WORD] Chart berhasil disisipkan")
    else:
        print("[WORD] Chart tidak ditemukan")

    # ---------------- SIMPAN FILE ----------------
    output_path = "laporan_cmvrp_profit.docx"
    doc.save(output_path)
    return send_file(output_path, as_attachment=True)

@app.route('/simpan_chart', methods=['POST'])
def simpan_chart():
    data_url = request.form.get('image')
    if not data_url:
        return "No image received", 400

    # Decode base64 dan simpan sebagai chart.png
    header, encoded = data_url.split(",", 1)
    image_data = b64decode(encoded)

    os.makedirs("static", exist_ok=True)
    with open("static/chart.png", "wb") as f:
        f.write(image_data)

    print("[CHART] Disimpan sebagai static/chart.png")
    return "OK"

@app.route('/laporan')
def tampilkan_laporan():
    if not os.path.exists("laporan_cmvrp.csv") or not os.path.exists("data_route_results.json"):
        return "Silakan klik tombol Proses dulu."

    df = pd.read_csv("laporan_cmvrp.csv")

    with open("data_route_results.json", "r") as f:
        route_data = json.load(f)

    # Helper: Normalisasi koordinat
    def get_coord(p):
        if isinstance(p, list) and isinstance(p[0], str):
            return tuple(p[1])
        elif isinstance(p, list) and isinstance(p[0], (float, int)):
            return tuple(p)
        else:
            return tuple(p)

    vehicle_routes = {
        str(r["vehicle_id"]): r.get("route_coords", [])
        for r in route_data
    }

    kendaraan_data = []
    total_jarak = total_fuel = total_ops = 0
    total_revenue = df["Revenue"].sum()
    total_biaya = 0
    profit_total = 0

    for _, row in df.iterrows():
        kendaraan_str = str(row["Kendaraan"]).replace("Kendaraan ", "").strip()
        coords = vehicle_routes.get(kendaraan_str, [])

        jarak_osrm = 0
        if len(coords) < 2:
            print(f"[SKIP] Kendaraan {kendaraan_str} memiliki <2 titik.")
            continue

        for i in range(len(coords) - 1):
            try:
                coord1 = get_coord(coords[i])
                coord2 = get_coord(coords[i + 1])
                result = osrm_request_for_report(coord1, coord2)
                jarak_osrm += result["distance_km"]
            except Exception as e:
                print(f"[OSRM] Kendaraan {kendaraan_str} error: {e}")

        fuel_cost = (jarak_osrm / fuel_consumption) * fuel_price
        biaya = row["Biaya"]
        ops_cost = biaya - fuel_cost
        profit = row["Revenue"] - (fuel_cost + ops_cost)

        kendaraan_data.append({
            "kendaraan": row["Kendaraan"],
            "jarak": round(jarak_osrm, 2),
            "overnight": row.get("Overnight (x)", 0),
            "revenue": round(row["Revenue"]),
            "biaya": round(fuel_cost + ops_cost),
            "profit": round(profit)
        })

        total_jarak += jarak_osrm
        total_fuel += fuel_cost
        total_ops += ops_cost
        total_biaya += fuel_cost + ops_cost
        profit_total += profit

        print(f"[Kendaraan {kendaraan_str}] Jarak: {jarak_osrm:.2f} km, Fuel Cost: {fuel_cost:.2f}, Ops Cost: {ops_cost:.2f}")
    html_laporan = render_template("laporan.html",
                                    kendaraan_data=kendaraan_data,
                                    total_fuel=round(total_fuel),
                                    total_ops=round(total_ops),
                                    total_biaya=round(total_biaya),
                                    total_revenue=round(total_revenue),
                                    profit_total=round(profit_total))

    with open("templates/laporan_profit_output.html", "w", encoding="utf-8") as f:
        f.write(html_laporan)

    return render_template("laporan.html",
        kendaraan_data=kendaraan_data,
        total_fuel=round(total_fuel),
        total_ops=round(total_ops),
        total_biaya=round(total_biaya),
        total_revenue=round(total_revenue),
        profit_total=round(profit_total),
        total_jarak=round(total_jarak)
    )

@app.route('/laporan_rute')
def tampilkan_laporan_rute():
    if not os.path.exists("data_route_results.json"):
        return "Silakan klik tombol Proses dulu."

    with open("data_route_results.json", "r") as f:
        route_data = json.load(f)

    _, customers, rest_areas, menginap_locs = load_input_data()

    hasil_laporan = []
    for vehicle in route_data:
        vehicle_id = vehicle["vehicle_id"]
        route_coords = vehicle["route_coords"]
        assigned_customers = vehicle["assigned_customers"]
        depot_nama, depot_data = vehicle["depot"]

        segmen = []
        total_distance = total_duration = total_pelayanan = total_istirahat = total_nginap = 0

        def get_coord(p):
            return tuple(p[1]) if isinstance(p, list) and isinstance(p[0], str) else tuple(p)

        for i in range(len(route_coords)):
            titik = route_coords[i]
            row = {"lokasi": "", "aktivitas": "", "waktu_servis": ""}

            coord = get_coord(titik)

            if isinstance(titik, list) and isinstance(titik[0], str):
                label = titik[0]
                if label == "rest_virtual":
                    row["lokasi"] = "Virtual Rest Area"
                    row["aktivitas"] = "Istirahat"
                    row["waktu_servis"] = "0.5 jam"
                    total_istirahat += 0.5
                elif label == "nginap_virtual":
                    row["lokasi"] = "Virtual Menginap"
                    row["aktivitas"] = "Menginap"
                    row["waktu_servis"] = "8 jam"
                    total_nginap += 8
                elif label == "customer":
                    match = next((c for c in assigned_customers if abs(c[0][0] - coord[0]) < 1e-5 and abs(c[0][1] - coord[1]) < 1e-5), None)
                    if match:
                        row["lokasi"] = match[1]
                        row["aktivitas"] = "Pelanggan"
                        row["waktu_servis"] = f"{match[2]:.2f} jam"
                        total_pelayanan += match[2]
            elif coord == tuple(depot_data[:2]):
                row["lokasi"] = depot_nama
                row["aktivitas"] = "Depot"
                row["waktu_servis"] = ""
            else:
                match = next((c for c in assigned_customers if abs(c[0][0] - coord[0]) < 1e-5 and abs(c[0][1] - coord[1]) < 1e-5), None)
                if match:
                    row["lokasi"] = match[1]
                    row["aktivitas"] = "Pelanggan"
                    row["waktu_servis"] = f"{match[2]:.2f} jam"
                    total_pelayanan += match[2]

            # Hitung waktu antar titik
            if i > 0:
                prev = get_coord(route_coords[i - 1])
                osrm_result = osrm_request_for_report(prev, coord)
                total_duration += osrm_result["duration_hr"]
                total_distance += osrm_result["distance_km"]

            if row["lokasi"]:
                segmen.append(row)

        total_work = total_duration + total_pelayanan + total_istirahat + total_nginap

        hasil_laporan.append({
            "kendaraan": f"Kendaraan {vehicle_id}",
            "rute": segmen,
            "kosong": len(segmen) == 0,
            "waktu_pelayanan": round(total_pelayanan, 2),
            "waktu_istirahat": total_istirahat,
            "waktu_nginap": total_nginap,
            "waktu_perjalanan": round(total_duration, 2),
            "total_work_time": round(total_work, 2)
        })
    html_rute = render_template("laporan_rute.html", rute_kendaraan=hasil_laporan)

    with open("templates/laporan_rute_output.html", "w", encoding="utf-8") as f:
        f.write(html_rute)

    return render_template("laporan_rute.html", rute_kendaraan=hasil_laporan)


@app.route('/proses')
def proses():
    for file in ["data_route_results.json", "vrp_tabu_search_visualisasi.html", "laporan_cmvrp.csv"]:
        if os.path.exists(file):
            os.remove(file)
    # Jalankan proses optimasi dan simpan hasil
    depots, customers, rest_areas, menginap_locs = load_input_data()
    vehicle_count = load_vehicle_count()
    assignments = allocate_customers(customers, depots, vehicle_count)
    # Cek pelanggan yang tidak terlayani
    unassigned_customers = set(customers.keys())
    for vehicle in assignments:
        for cust in vehicle["customers"]:
            unassigned_customers.discard(cust)

    if unassigned_customers:
        flash(f"{len(unassigned_customers)} pelanggan tidak terlayani. Tambahkan kendaraan atau kurangi permintaan.", "danger")
    
    all_route_results = []
    laporan_csv = []

    for i, vehicle in enumerate(assignments, start=1):
        customers_list = vehicle.get("customers", [])
        depot_info = vehicle.get("depot", ("Unknown", [0.0, 0.0]))

        if not customers_list:
            continue

        depot_coord = depot_info[1]
        route_result = tabu_search_vrp(
            depot_coord, customers_list, customers,
            rest_areas=rest_areas, menginap_locs=menginap_locs
        )
        if not route_result:
            continue
        route_coords, assigned_customers, overnight, dist, revenue, cost, profit, service_time, rest_time, nginap_time, total_time, travel_time, work_time = route_result

        # Simpan data hasil untuk visualisasi dan ekspor
        all_route_results.append({
            "vehicle_id": i,
            "route_coords": route_coords,
            "assigned_customers": assigned_customers,
            "depot": depot_info,
            "overnight_stays": overnight,
            "total_distance": dist,
            "total_revenue": revenue,
            "total_cost": cost,
            "profit": profit,
            "service_time": service_time,
            "rest_time": rest_time,
            "nginap_time": nginap_time,
            "total_time": total_time,
            "travel_time": travel_time,
            "work_time": work_time
        })

        laporan_csv.append({
            "Kendaraan": f"Kendaraan {i}",
            "Jarak Tempuh (km)": round(dist, 2),
            "Overnight (x)": overnight,
            "Waktu Pelayanan (jam)": round(service_time, 2),
            "Waktu Istirahat (jam)": round(rest_time, 2),
            "Waktu Menginap (jam)": round(nginap_time, 2),
            "Waktu Perjalanan (jam)": round(travel_time, 2),
            "Total Waktu (jam)": round(total_time, 2),
            "Biaya": round(cost),
            "Revenue": round(revenue),
            "Profit": round(profit)
        })

    if not all_route_results:
        flash("Gagal memproses rute: jumlah kendaraan tidak mencukupi untuk melayani semua pelanggan.", "warning")
        return redirect("/")
    # Simpan hasil untuk ekspor dan tampilan kembali
    with open("data_route_results.json", "w") as f:
        json.dump(all_route_results, f, indent=4)
    pd.DataFrame(laporan_csv).to_csv("laporan_cmvrp.csv", index=False)

    print("ISI ALL ROUTE RESULTS:", all_route_results)
    print("TIPE:", type(all_route_results), [type(item) for item in all_route_results])

    # ✅ Konversi ke format tuple agar cocok untuk visualisasi
    converted_routes = []
    for r in all_route_results:
        converted_routes.append((
            r["route_coords"],
            r["assigned_customers"],
            r.get("overnight_stays", 0),
            r.get("total_distance", 0),
            r.get("total_revenue", 0),
            r.get("total_cost", 0),
            r.get("profit", 0),
            r.get("service_time", 0),
            r.get("rest_time", 0),
            r.get("nginap_time", 0),
            r.get("total_time", 0),
            r.get("travel_time", 0),
            r.get("work_time", 0)
        ))

    # Kirim ke fungsi visualisasi
    generate_leaflet_html(converted_routes, rest_areas, menginap_locs)
    with app.test_request_context():
        tampilkan_laporan()
        tampilkan_laporan_rute()
    flash("Proses optimasi selesai! Silakan cek laporan rute kendaraan.", "success")
    return redirect('/laporan_rute')

@app.route('/hapus_hasil')
def hapus_hasil():
    files_to_delete = [
        "data_route_results.json",
        "vrp_tabu_search_visualisasi.html",
        "laporan_cmvrp.csv",
        "laporan_segment_cmvrp.csv",
        "templates/laporan_rute_output.html",
        "templates/laporan_profit_output.html", "osrm_cache.json"
    ]
    
    deleted_files = []
    for file in files_to_delete:
        try:
            if os.path.exists(file):
                os.remove(file)
                deleted_files.append(file)
        except Exception as e:
            print(f"[HAPUS ERROR] Tidak bisa hapus {file}: {e}")

    flash(f"Hasil proses berhasil dihapus ({len(deleted_files)} file).", "info")
    return redirect('/')

if __name__ == '__main__':
    start_osrm()
    app.run(debug=True)